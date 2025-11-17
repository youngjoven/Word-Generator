"""
Word 산출물 문서 자동 생성기

YAML 파일로부터 Word 문서(.docx)를 자동 생성합니다.
- 표지 페이지 (제목, 프로젝트 정보)
- 목차 (하이퍼링크, 점선 리더)
- 본문 (계층적 섹션 구조)
- 페이지 번호 자동 삽입
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, Any, List


class DocumentGenerator:
    """Word 산출물 문서를 생성하는 클래스"""

    def __init__(self):
        self.doc = Document()
        self._setup_document()
        self._setup_styles()
        self._enable_auto_update_fields()

    def _setup_document(self):
        """문서 기본 설정"""
        section = self.doc.sections[0]

        # A4 용지 크기
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

        # 여백 설정
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.18)
        section.bottom_margin = Inches(2.07)

    def _setup_styles(self):
        """문서 스타일 설정"""
        # Normal 스타일
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '맑은 고딕'
        font.size = Pt(10)

        # Heading 스타일
        heading_sizes = {1: 16, 2: 13, 3: 12, 4: 11, 5: 10}
        for i in range(1, 6):
            heading_style = self.doc.styles[f'Heading {i}']
            heading_font = heading_style.font
            heading_font.name = '맑은 고딕'
            heading_font.bold = True
            heading_font.size = Pt(heading_sizes.get(i, 10))

    def _enable_auto_update_fields(self):
        """문서를 열 때 자동으로 필드가 업데이트되도록 설정"""
        # settings.xml에 updateFields 설정 추가
        settings_element = self.doc.settings.element

        # updateFields 요소 생성
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')

        # settings에 추가
        settings_element.append(update_fields)

    def _add_page_number_footer(self):
        """푸터에 페이지 번호 추가"""
        section = self.doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 페이지 번호 필드
        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

        paragraph.add_run(' / ')

        # 총 페이지 수 필드
        run2 = paragraph.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = 'NUMPAGES'
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run2._r.append(fldChar3)
        run2._r.append(instrText2)
        run2._r.append(fldChar4)

    def generate_document(self, data: Dict[str, Any], output_path: str):
        """
        완전한 Word 문서 생성

        Args:
            data: 문서 데이터 (metadata, sections)
            output_path: 출력 파일 경로
        """
        # 1. 표지 페이지
        self._add_cover_page(data.get('metadata', {}))
        self.doc.add_page_break()

        # 2. 섹션 정보 수집 (목차용)
        sections = data.get('sections', [])
        self.toc_entries = []
        self._collect_toc_entries(sections, level=1)

        # 3. 목차 페이지 (표 형태)
        self._add_table_of_contents()
        self.doc.add_page_break()

        # 4. 푸터에 페이지 번호 추가
        self._add_page_number_footer()

        # 5. 본문 섹션
        for idx, section in enumerate(sections, 1):
            # 첫 번째 섹션이 아니면 페이지 나누기
            if idx > 1:
                self.doc.add_page_break()

            # 섹션 추가 (Heading 1마다 새 페이지)
            self._add_section(section, level=1, parent_num="", section_idx=idx)

        # 문서 저장
        self.doc.save(output_path)

    def _add_cover_page(self, metadata: Dict[str, Any]):
        """표지 페이지 생성"""
        # 빈 줄 (상단 여백)
        for _ in range(8):
            self.doc.add_paragraph()

        # 제목 - 크게, 중앙 정렬, 굵게
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(metadata.get('title', 'Azure Infrastructure 구성 가이드'))
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = '맑은 고딕'

        # 빈 줄
        for _ in range(10):
            self.doc.add_paragraph()

        # 프로젝트 정보 표
        table = self.doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        # 표 내용
        info_fields = [
            ('프로젝트명', metadata.get('project_name', '')),
            ('고객사', metadata.get('client', '')),
            ('작성일자', metadata.get('date', '')),
            ('작성자', metadata.get('author', '')),
            ('버전', metadata.get('version', '1.0'))
        ]

        for row_idx, (label, value) in enumerate(info_fields):
            row = table.rows[row_idx]
            row.cells[0].text = label
            row.cells[1].text = value

            # 첫 번째 열 굵게
            for para in row.cells[0].paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(11)

            # 두 번째 열 스타일
            for para in row.cells[1].paragraphs:
                for run in para.runs:
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(11)

    def _collect_toc_entries(self, sections: List[Dict[str, Any]], level: int, parent_num: str = ""):
        """목차 항목을 재귀적으로 수집"""
        for idx, section in enumerate(sections, 1):
            # 번호 생성
            if parent_num:
                number = f"{parent_num}.{idx}"
            else:
                number = str(idx)

            title = section.get('title', '')

            # 목차 항목 추가 (제목, 레벨, 북마크명)
            bookmark_name = f"heading_{number.replace('.', '_')}"
            self.toc_entries.append({
                'title': title,
                'level': level,
                'bookmark': bookmark_name,
                'number': number
            })

            # 하위 섹션 재귀 처리
            subsections = section.get('subsections', [])
            if subsections:
                self._collect_toc_entries(subsections, level + 1, number)

    def _add_table_of_contents(self):
        """점선으로 연결된 목차 생성 (하이퍼링크 있음, 밑줄 없음)"""
        # 목차 제목
        self.doc.add_heading('목차', level=1)
        self.doc.add_paragraph()

        # 목차 항목들을 단락으로 추가
        for entry in self.toc_entries:
            para = self.doc.add_paragraph()

            # 들여쓰기 설정 (레벨에 따라)
            indent_level = entry['level'] - 1
            para.paragraph_format.left_indent = Inches(0.25 * indent_level)

            # 탭 설정 (점선 리더)
            tab_stops = para.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.0), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

            # 제목 (하이퍼링크 있음, 밑줄 없음)
            self._add_hyperlink_no_underline(para, entry['bookmark'], f"{entry['number']}. {entry['title']}")

            # 탭 추가
            para.add_run('\t')

            # 페이지 번호 (PAGEREF 필드)
            run = para.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')

            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = f'PAGEREF {entry["bookmark"]} \\h'

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)

            # 폰트 설정
            for run in para.runs:
                run.font.name = '맑은 고딕'
                run.font.size = Pt(10)

        # 안내 메시지
        self.doc.add_paragraph()
        note = self.doc.add_paragraph()
        note_run = note.add_run('※ Word에서 문서를 열어 목차 영역 전체를 선택 → 우클릭 → "필드 업데이트"를 하면 페이지 번호가 자동으로 업데이트됩니다.')
        note_run.font.size = Pt(9)
        note_run.font.italic = True
        note_run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_hyperlink_no_underline(self, paragraph, bookmark_name, text):
        """단락에 하이퍼링크 추가 (밑줄 없음, 검은색)"""
        # 하이퍼링크 요소 생성
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)

        # 런 생성
        run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # 스타일 (검은색, 밑줄 없음)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), '000000')  # 검은색
        rPr.append(color)

        # 밑줄 제거
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'none')
        rPr.append(u)

        run.append(rPr)

        # 텍스트 추가
        t = OxmlElement('w:t')
        t.text = text
        run.append(t)

        hyperlink.append(run)
        paragraph._p.append(hyperlink)


    def _add_section(self, section: Dict[str, Any], level: int = 1, parent_num: str = "", section_idx: int = 1):
        """
        섹션 추가 (재귀적)

        Args:
            section: 섹션 데이터
            level: 제목 레벨
            parent_num: 부모 번호
            section_idx: 섹션 인덱스
        """
        # 번호 생성
        if parent_num:
            number = f"{parent_num}.{section_idx}"
        else:
            number = str(section_idx)

        # 북마크명
        bookmark_name = f"heading_{number.replace('.', '_')}"

        # 제목
        title = section.get('title', '')
        heading = self.doc.add_heading(title, level=level)

        # 북마크 추가
        self._add_bookmark(heading, bookmark_name)

        # 내용 (content)
        content = section.get('content', '')
        if content:
            for line in content.split('\n'):
                line = line.strip()
                if line:
                    para = self.doc.add_paragraph(line)
                    try:
                        para.style = 'List Paragraph'
                    except:
                        pass

        # 단계별 설명 (steps)
        steps = section.get('steps', [])
        for step in steps:
            para = self.doc.add_paragraph(step)
            try:
                para.style = 'List Paragraph'
            except:
                pass

        # 하위 섹션
        subsections = section.get('subsections', [])
        for idx, subsection in enumerate(subsections, 1):
            self._add_section(subsection, level + 1, number, idx)

    def _add_bookmark(self, paragraph, bookmark_name):
        """단락에 북마크 추가"""
        # 북마크 시작
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')
        bookmark_start.set(qn('w:name'), bookmark_name)

        # 북마크 끝
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')

        # 단락에 북마크 추가
        paragraph._p.insert(0, bookmark_start)
        paragraph._p.append(bookmark_end)


